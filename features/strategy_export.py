# features/strategy_export.py
import streamlit as st
import os  # ✅ إضافة استيراد os
import os
from io import BytesIO
from core.i18n import get_text

def export_to_word(strategy_data):
    """تصدير الاستراتيجية إلى ملف Word"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from PIL import Image
    except ImportError:
        st.error("مكتبة python-docx غير مثبتة. يرجى تثبيتها لتمكين تصدير Word.")
        return

    # إعدادات الخط العربي
    ARABIC_FONTS = [
        "Sakkal Majalla",
        "Traditional Arabic",
        "Arabic Typesetting",
        "Arial",
        "Times New Roman"
    ]
    
    AR_FONT_FAMILY = "Sakkal Majalla"

    def _run_set_ar_font(run, bold=False, font_size=14):
        try:
            for font in ARABIC_FONTS:
                try:
                    run.font.name = font
                    break
                except:
                    continue
            
            run.font.bold = bool(bold)
            run.font.size = Pt(font_size)
            
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), AR_FONT_FAMILY)
            rFonts.set(qn('w:hAnsi'), AR_FONT_FAMILY)
            rFonts.set(qn('w:cs'), AR_FONT_FAMILY)
            lang = rPr.find(qn('w:lang'))
            if lang is None:
                lang = OxmlElement('w:lang')
                rPr.append(lang)
            lang.set(qn('w:bidi'), 'ar-SA')
            lang.set(qn('w:eastAsia'), 'ar-SA')
        except Exception:
            pass

    def _rtl_paragraph(p):
        try:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Inches(0.08)
            p.paragraph_format.space_before = Inches(0.04)
            p.paragraph_format.line_spacing = 1.5
            
            pPr = p._p.get_or_add_pPr()
            bidi = pPr.find(qn('w:bidi'))
            if bidi is None:
                bidi = OxmlElement('w:bidi')
                pPr.append(bidi)
            bidi.set(qn('w:val'), '1')
            
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                pPr.append(jc)
            jc.set(qn('w:val'), 'both')
            
            for r in p.runs:
                _run_set_ar_font(r)
        except Exception:
            pass

    def _add_heading_rtl(doc, text, level=1):
        style = "Heading 1" if level == 1 else "Heading 2"
        p = doc.add_paragraph("", style=style)
        _rtl_paragraph(p)
        r = p.add_run(text)
        _run_set_ar_font(r, bold=True, font_size=16)
        return p

    def _add_paragraph_rtl(doc, text, style=None):
        p = doc.add_paragraph("", style=style) if style else doc.add_paragraph("")
        _rtl_paragraph(p)
        r = p.add_run(text)
        _run_set_ar_font(r, bold=False, font_size=14)
        return p

    def _add_bullet_rtl(doc, text):
        p = doc.add_paragraph("", style="List Bullet")
        _rtl_paragraph(p)
        r = p.add_run(text)
        _run_set_ar_font(r, bold=False, font_size=14)
        
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        return p

    def _ensure_doc_defaults_rtl(doc):
        try:
            styles_elm = doc.styles.element
            docDefaults = styles_elm.find(qn('w:docDefaults'))
            if docDefaults is None:
                docDefaults = OxmlElement('w:docDefaults')
                styles_elm.append(docDefaults)
            pPrDefault = docDefaults.find(qn('w:pPrDefault'))
            if pPrDefault is None:
                pPrDefault = OxmlElement('w:pPrDefault')
                docDefaults.append(pPrDefault)
            pPr = pPrDefault.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                pPrDefault.append(pPr)
            bidi = pPr.find(qn('w:bidi'))
            if bidi is None:
                bidi = OxmlElement('w:bidi')
                pPr.append(bidi)
            bidi.set(qn('w:val'), '1')
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                pPr.append(jc)
            jc.set(qn('w:val'), 'both')
            
            rPrDefault = docDefaults.find(qn('w:rPrDefault'))
            if rPrDefault is None:
                rPrDefault = OxmlElement('w:rPrDefault')
                docDefaults.append(rPrDefault)
            rPr = rPrDefault.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                rPrDefault.append(rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), AR_FONT_FAMILY)
            rFonts.set(qn('w:hAnsi'), AR_FONT_FAMILY)
            rFonts.set(qn('w:cs'), AR_FONT_FAMILY)
            
        except Exception:
            pass

    try:
        doc = Document()
        _ensure_doc_defaults_rtl(doc)
        
        for st_name in ["Normal", "Heading 1", "Heading 2", "List Bullet"]:
            try:
                style = doc.styles[st_name]
                style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                style.paragraph_format.space_after = Inches(0.08)
                style.paragraph_format.space_before = Inches(0.04)
                style.paragraph_format.line_spacing = 1.5
            except Exception:
                pass

        vision = strategy_data.get("edited_vision", strategy_data.get("vision", ""))
        message = strategy_data.get("edited_message", strategy_data.get("message", ""))

        logo_path = strategy_data.get("logo_path")
        if logo_path and os.path.exists(logo_path):
            try:
                with Image.open(logo_path) as img:
                    width, height = img.size
                
                max_width = 4.0
                max_height = 3.0
                
                ratio = min(max_width/width, max_height/height)
                new_width = width * ratio
                new_height = height * ratio
                
                logo_paragraph = doc.add_paragraph()
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                run = logo_paragraph.add_run()
                run.add_picture(logo_path, width=Inches(new_width))
                
                doc.add_paragraph()
                
            except Exception as e:
                print(f"⚠️ Could not add logo to Word: {e}")
                try:
                    logo_paragraph = doc.add_paragraph()
                    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = logo_paragraph.add_run()
                    run.add_picture(logo_path, width=Inches(3.0))
                    doc.add_paragraph()
                except:
                    pass

        _add_heading_rtl(doc, str(strategy_data.get("name", "استراتيجية")), level=1)
        doc.add_paragraph()
        
        if strategy_data.get("description"):
            _add_heading_rtl(doc, "الوصف", level=2)
            _add_paragraph_rtl(doc, str(strategy_data.get("description", "")))
            doc.add_paragraph()
        
        _add_heading_rtl(doc, "الرؤية", level=2)
        _add_paragraph_rtl(doc, str(vision))
        doc.add_paragraph()
        
        _add_heading_rtl(doc, "الرسالة", level=2)
        _add_paragraph_rtl(doc, str(message))
        doc.add_paragraph()
        
        _add_heading_rtl(doc, "الأهداف", level=2)
        for obj in strategy_data.get("objectives", []):
            _add_bullet_rtl(doc, str(obj))
        doc.add_paragraph()
        
        _add_heading_rtl(doc, "القيم", level=2)
        for val in strategy_data.get("values", []):
            _add_bullet_rtl(doc, str(val))

        for p in doc.paragraphs:
            _rtl_paragraph(p)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        st.success("✅ تم إنشاء ملف Word بنجاح!")
        st.download_button(
            label=get_text("export_to_word"),
            data=buffer,
            file_name=f"{strategy_data['name'].replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        st.error(f"خطأ في إنشاء ملف Word: {str(e)}")

def export_to_pdf(strategy_data):
    """تصدير الاستراتيجية إلى ملف PDF"""
    try:
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_RIGHT, TA_CENTER
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfbase import pdfmetrics
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import inch
        import arabic_reshaper
        from bidi.algorithm import get_display
        from PIL import Image as PILImage
    except ImportError:
        st.error("المكتبات المطلوبة غير مثبتة. يرجى تثبيت: pip install reportlab arabic-reshaper python-bidi")
        return

    def _process_arabic_text(text):
        try:
            if not text:
                return ""
            reshaped_text = arabic_reshaper.reshape(str(text))
            bidi_text = get_display(reshaped_text)
            return bidi_text
        except:
            return str(text)

    def _register_noto_font():
        try:
            noto_paths = [
                "fonts/NotoNaskhArabic-Regular.ttf",
                "assets/fonts/NotoNaskhArabic-Regular.ttf",
                "NotoNaskhArabic-Regular.ttf",
                "C:/Windows/Fonts/NotoNaskhArabic-Regular.ttf",
                "./NotoNaskhArabic-Regular.ttf",
            ]
            
            for font_path in noto_paths:
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont('NotoNaskhArabic', font_path))
                    
                    bold_path = font_path.replace('Regular', 'Bold')
                    if os.path.exists(bold_path):
                        pdfmetrics.registerFont(TTFont('NotoNaskhArabic-Bold', bold_path))
                        return 'NotoNaskhArabic', 'NotoNaskhArabic-Bold'
                    else:
                        pdfmetrics.registerFont(TTFont('NotoNaskhArabic-Bold', font_path))
                        return 'NotoNaskhArabic', 'NotoNaskhArabic-Bold'
            
            st.warning("⚠️ خط Noto Naskh Arabic غير موجود. جاري البحث عن خطوط بديلة...")
            
            arabic_fonts = [
                "arial.ttf", "tahoma.ttf", "times.ttf",
                "C:/Windows/Fonts/arial.ttf",
                "C:/Windows/Fonts/tahoma.ttf", 
            ]
            
            for font_path in arabic_fonts:
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont('ArabicFont', font_path))
                    pdfmetrics.registerFont(TTFont('ArabicFont-Bold', font_path))
                    st.warning(f"⚠️ استخدام خط بديل: {os.path.basename(font_path)}")
                    return 'ArabicFont', 'ArabicFont-Bold'
            
            st.error("❌ لم يتم العثور على أي خط يدعم العربية")
            return 'Helvetica', 'Helvetica-Bold'
            
        except Exception as e:
            st.error(f"❌ خطأ في تحميل الخط: {e}")
            return 'Helvetica', 'Helvetica-Bold'

    buffer = BytesIO()
    
    normal_font, bold_font = _register_noto_font()
    
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=A4,
        rightMargin=40, 
        leftMargin=40, 
        topMargin=40, 
        bottomMargin=40,
        encoding='utf-8'
    )
    
    vision = strategy_data.get("edited_vision", strategy_data.get("vision", ""))
    message = strategy_data.get("edited_message", strategy_data.get("message", ""))
    
    styles = getSampleStyleSheet()
    
    style_title = ParagraphStyle(
        'ArabicTitle',
        parent=styles['Title'],
        fontName=bold_font,
        alignment=TA_CENTER,
        fontSize=20,
        leading=28,
        spaceAfter=20,
        wordWrap='RTL',
        textColor='#2c3e50'
    )
    
    style_heading = ParagraphStyle(
        'ArabicHeading',
        parent=styles['Heading1'],
        fontName=bold_font,
        alignment=TA_RIGHT,
        fontSize=16,
        spaceBefore=18,
        spaceAfter=10,
        wordWrap='RTL',
        textColor='#2980b9'
    )
    
    style_normal = ParagraphStyle(
        'ArabicNormal',
        parent=styles['Normal'],
        fontName=normal_font,
        alignment=TA_RIGHT,
        fontSize=14,
        leading=20,
        wordWrap='RTL',
        textColor='#2c3e50'
    )
    
    style_bullet = ParagraphStyle(
        'ArabicBullet',
        parent=styles['Normal'],
        fontName=normal_font,
        alignment=TA_RIGHT,
        fontSize=14,
        leading=18,
        leftIndent=30,
        wordWrap='RTL',
        textColor='#2c3e50'
    )
    
    story = []
    
    logo_path = strategy_data.get("logo_path")
    if logo_path and os.path.exists(logo_path):
        try:
            with PILImage.open(logo_path) as img:
                width, height = img.size
            
            max_width = 3.0 * inch
            max_height = 2.5 * inch
            
            ratio = min(max_width/width, max_height/height)
            new_width = width * ratio
            new_height = height * ratio
            
            logo = Image(logo_path, width=new_width, height=new_height)
            logo.hAlign = 'CENTER'
            story.append(logo)
            story.append(Spacer(1, 20))
            
        except Exception as e:
            print(f"⚠️ Could not add logo to PDF with natural size: {e}")
            try:
                logo = Image(logo_path, width=2.5*inch, height=2.0*inch)
                logo.hAlign = 'CENTER'
                story.append(logo)
                story.append(Spacer(1, 15))
            except Exception as e2:
                print(f"⚠️ Could not add logo to PDF: {e2}")

    title_text = _process_arabic_text(strategy_data["name"])
    if title_text:
        story.append(Paragraph(title_text, style_title))
        story.append(Spacer(1, 20))
    
    if strategy_data.get("description"):
        desc_text = _process_arabic_text(strategy_data["description"])
        if desc_text:
            story.append(Paragraph(desc_text, style_normal))
            story.append(Spacer(1, 15))
    
    vision_heading = _process_arabic_text(get_text("select_vision"))
    vision_text = _process_arabic_text(vision)
    if vision_heading and vision_text:
        story.append(Paragraph(vision_heading, style_heading))
        story.append(Paragraph(vision_text, style_normal))
        story.append(Spacer(1, 15))
    
    message_heading = _process_arabic_text(get_text("select_message"))
    message_text = _process_arabic_text(message)
    if message_heading and message_text:
        story.append(Paragraph(message_heading, style_heading))
        story.append(Paragraph(message_text, style_normal))
        story.append(Spacer(1, 15))
    
    objectives_heading = _process_arabic_text(get_text("strategic_objectives"))
    if objectives_heading and strategy_data.get("objectives"):
        story.append(Paragraph(objectives_heading, style_heading))
        for obj in strategy_data["objectives"]:
            obj_text = _process_arabic_text("• " + str(obj))
            if obj_text:
                story.append(Paragraph(obj_text, style_bullet))
        story.append(Spacer(1, 15))
    
    values_heading = _process_arabic_text(get_text("strategic_values"))
    if values_heading and strategy_data.get("values"):
        story.append(Paragraph(values_heading, style_heading))
        for val in strategy_data["values"]:
            val_text = _process_arabic_text("• " + str(val))
            if val_text:
                story.append(Paragraph(val_text, style_bullet))
    
    try:
        if story:
            doc.build(story)
            buffer.seek(0)
            
            st.success("✅ تم إنشاء ملف PDF بنجاح باستخدام خط Noto Naskh Arabic!")
            st.download_button(
                label=get_text("export_to_pdf"),
                data=buffer,
                file_name=f"{strategy_data['name'].replace(' ', '_')}.pdf",
                mime="application/pdf"
            )
        else:
            st.error("لا يوجد محتوى لتصديره إلى PDF")
    except Exception as e:
        st.error(f"خطأ في إنشاء PDF: {str(e)}")
        try:
            st.info("جاري استخدام نسخة مبسطة...")
            export_to_pdf_simple(strategy_data)
        except:
            st.error("فشل إنشاء PDF. يرجى التحقق من تثبيت المكتبات المطلوبة.")

def export_to_pdf_simple(strategy_data):
    """نسخة مبسطة للطوارئ لتصدير PDF"""
    try:
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_RIGHT, TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import inch
    except ImportError:
        st.error("مكتبة reportlab غير مثبتة.")
        return

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    
    style_normal = ParagraphStyle(
        'SimpleArabic',
        parent=styles['Normal'],
        alignment=TA_RIGHT,
        fontSize=12,
        leading=16
    )
    
    style_heading = ParagraphStyle(
        'SimpleArabicHeading', 
        parent=styles['Heading2'],
        alignment=TA_RIGHT,
        fontSize=14,
        spaceAfter=8
    )
    
    story = []
    
    logo_path = strategy_data.get("logo_path")
    if logo_path and os.path.exists(logo_path):
        try:
            logo = Image(logo_path, width=2.0*inch, height=1.5*inch)
            logo.hAlign = 'CENTER'
            story.append(logo)
            story.append(Spacer(1, 15))
        except:
            pass
    
    vision = strategy_data.get("edited_vision", strategy_data.get("vision", ""))
    message = strategy_data.get("edited_message", strategy_data.get("message", ""))
    
    story.append(Paragraph(str(strategy_data["name"]), styles['Title']))
    story.append(Spacer(1, 12))
    
    if strategy_data.get("description"):
        story.append(Paragraph(str(strategy_data["description"]), style_normal))
        story.append(Spacer(1, 12))
    
    story.append(Paragraph("الرؤية:", style_heading))
    story.append(Paragraph(str(vision), style_normal))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("الرسالة:", style_heading))
    story.append(Paragraph(str(message), style_normal))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("الأهداف:", style_heading))
    for obj in strategy_data.get("objectives", []):
        story.append(Paragraph("• " + str(obj), style_normal))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("القيم:", style_heading))
    for val in strategy_data.get("values", []):
        story.append(Paragraph("• " + str(val), style_normal))
    
    doc.build(story)
    buffer.seek(0)
    
    st.download_button(
        label=get_text("export_to_pdf"),
        data=buffer,
        file_name=f"{strategy_data['name'].replace(' ', '_')}.pdf",
        mime="application/pdf"
    )