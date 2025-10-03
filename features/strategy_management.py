# features/strategy_management.py - النسخة النهائية المصححة
import streamlit as st
import os
from io import BytesIO
from core.i18n import get_text, LANGS
from users.auth_service import get_current_user
from data.models import (
    create_strategy,
    get_user_strategies,
    get_strategy_by_id,
    update_strategy,
    delete_strategy,
    save_custom_element
)
from data.strategy_elements import (
    get_all_visions,
    get_messages_for_vision,
    get_all_objectives,
    get_objectives_for_vision_and_message,
    get_all_values,
    get_values_for_vision_and_message
)
from core.file_utils import validate_logo_file, save_logo_file, delete_logo_file, get_logo_url

# --- دوال التصدير المحسنة ---

def export_to_word(strategy_data):
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import os
        from PIL import Image  # ✅ لإضافة حجم الصورة الطبيعي
    except ImportError:
        st.error("مكتبة python-docx غير مثبتة. يرجى تثبيتها لتمكين تصدير Word.")
        return

    # إعدادات الخط العربي - استخدام Sakkal Majalla
    ARABIC_FONTS = [
        "Sakkal Majalla",      # الخط الأساسي المطلوب
        "Traditional Arabic",  # بديل أول
        "Arabic Typesetting",  # بديل ثاني
        "Arial",               # بديل يدعم العربية
        "Times New Roman"      # بديل كلاسيكي
    ]
    
    AR_FONT_FAMILY = "Sakkal Majalla"  # الخط الافتراضي المطلوب

    def _run_set_ar_font(run, bold=False, font_size=14):
        try:
            # محاولة استخدام الخطوط العربية بالترتيب
            for font in ARABIC_FONTS:
                try:
                    run.font.name = font
                    break
                except:
                    continue
            
            run.font.bold = bool(bold)
            run.font.size = Pt(font_size)  # تحديد حجم الخط
            
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), AR_FONT_FAMILY)
            rFonts.set(qn('w:hAnsi'), AR_FONT_FAMILY)
            rFonts.set(qn('w:cs'), AR_FONT_FAMILY)
            lang = rPr.find(qn('w:lang'))
            if lang is None:
                lang = OxmlElement('w:lang')
                rPr.append(lang)
            lang.set(qn('w:bidi'), 'ar-SA')
            lang.set(qn('w:eastAsia'), 'ar-SA')
        except Exception:
            pass

    def _rtl_paragraph(p):
        try:
            # محاذاة justify للنص العربي
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Inches(0.08)
            p.paragraph_format.space_before = Inches(0.04)
            p.paragraph_format.line_spacing = 1.5
            
            pPr = p._p.get_or_add_pPr()
            bidi = pPr.find(qn('w:bidi'))
            if bidi is None:
                bidi = OxmlElement('w:bidi')
                pPr.append(bidi)
            bidi.set(qn('w:val'), '1')
            
            # إعدادات justify إضافية
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                pPr.append(jc)
            jc.set(qn('w:val'), 'both')  # both = justify
            
            for r in p.runs:
                _run_set_ar_font(r)
        except Exception:
            pass

    def _add_heading_rtl(doc, text, level=1):
        style = "Heading 1" if level == 1 else "Heading 2"
        p = doc.add_paragraph("", style=style)
        _rtl_paragraph(p)
        r = p.add_run(text)
        _run_set_ar_font(r, bold=True, font_size=16)  # حجم 16 للعناوين
        return p

    def _add_paragraph_rtl(doc, text, style=None):
        p = doc.add_paragraph("", style=style) if style else doc.add_paragraph("")
        _rtl_paragraph(p)
        r = p.add_run(text)
        _run_set_ar_font(r, bold=False, font_size=14)  # حجم 14 للكتابة العادية
        return p

    def _add_bullet_rtl(doc, text):
        p = doc.add_paragraph("", style="List Bullet")
        _rtl_paragraph(p)
        r = p.add_run(text)
        _run_set_ar_font(r, bold=False, font_size=14)  # حجم 14 للنقاط
        
        # تحسين المسافات للنقاط
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        return p

    def _ensure_doc_defaults_rtl(doc):
        try:
            styles_elm = doc.styles.element
            docDefaults = styles_elm.find(qn('w:docDefaults'))
            if docDefaults is None:
                docDefaults = OxmlElement('w:docDefaults')
                styles_elm.append(docDefaults)
            pPrDefault = docDefaults.find(qn('w:pPrDefault'))
            if pPrDefault is None:
                pPrDefault = OxmlElement('w:pPrDefault')
                docDefaults.append(pPrDefault)
            pPr = pPrDefault.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                pPrDefault.append(pPr)
            bidi = pPr.find(qn('w:bidi'))
            if bidi is None:
                bidi = OxmlElement('w:bidi')
                pPr.append(bidi)
            bidi.set(qn('w:val'), '1')
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                pPr.append(jc)
            jc.set(qn('w:val'), 'both')  # justify للمستند كله
            
            # إعدادات الخط الافتراضي
            rPrDefault = docDefaults.find(qn('w:rPrDefault'))
            if rPrDefault is None:
                rPrDefault = OxmlElement('w:rPrDefault')
                docDefaults.append(rPrDefault)
            rPr = rPrDefault.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                rPrDefault.append(rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), AR_FONT_FAMILY)
            rFonts.set(qn('w:hAnsi'), AR_FONT_FAMILY)
            rFonts.set(qn('w:cs'), AR_FONT_FAMILY)
            
        except Exception:
            pass

    try:
        # إنشاء المستند
        doc = Document()
        _ensure_doc_defaults_rtl(doc)
        
        # ضبط الأنماط الأساسية للمستند
        for st_name in ["Normal", "Heading 1", "Heading 2", "List Bullet"]:
            try:
                style = doc.styles[st_name]
                style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                style.paragraph_format.space_after = Inches(0.08)
                style.paragraph_format.space_before = Inches(0.04)
                style.paragraph_format.line_spacing = 1.5
            except Exception:
                pass

        # استخدام القيم المحررة إذا كانت موجودة
        vision = strategy_data.get("edited_vision", strategy_data.get("vision", ""))
        message = strategy_data.get("edited_message", strategy_data.get("message", ""))

        # ✅ إضافة اللوجو بحجمه الطبيعي ومتوسطاً - مرة واحدة فقط
        logo_path = strategy_data.get("logo_path")
        if logo_path and os.path.exists(logo_path):
            try:
                # الحصول على الأبعاد الطبيعية للصورة
                with Image.open(logo_path) as img:
                    width, height = img.size
                
                # حساب الحجم المناسب مع الحفاظ على النسبة
                max_width = 4.0  # أقصى عرض 4 إنش
                max_height = 3.0  # أقصى ارتفاع 3 إنش
                
                ratio = min(max_width/width, max_height/height)
                new_width = width * ratio
                new_height = height * ratio
                
                # إضافة اللوجو في فقرة مخصصة بمحاذاة الوسط
                logo_paragraph = doc.add_paragraph()
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # إضافة الصورة بالحجم المحسوب
                run = logo_paragraph.add_run()
                run.add_picture(logo_path, width=Inches(new_width))
                
                doc.add_paragraph()  # سطر فارغ بعد اللوجو
                
            except Exception as e:
                print(f"⚠️ Could not add logo to Word: {e}")
                # محاولة بديلة بإضافة الصورة بحجم افتراضي
                try:
                    logo_paragraph = doc.add_paragraph()
                    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = logo_paragraph.add_run()
                    run.add_picture(logo_path, width=Inches(3.0))
                    doc.add_paragraph()
                except:
                    pass

        # العنوان الرئيسي
        _add_heading_rtl(doc, str(strategy_data.get("name", "استراتيجية")), level=1)
        doc.add_paragraph()  # سطر فارغ
        
        # الوصف إذا موجود
        if strategy_data.get("description"):
            _add_heading_rtl(doc, "الوصف", level=2)
            _add_paragraph_rtl(doc, str(strategy_data.get("description", "")))
            doc.add_paragraph()  # سطر فارغ
        
        # الرؤية - استخدام "الرؤية" بدلاً من "اختر رؤية"
        _add_heading_rtl(doc, "الرؤية", level=2)
        _add_paragraph_rtl(doc, str(vision))
        doc.add_paragraph()  # سطر فارغ
        
        # الرسالة - استخدام "الرسالة" بدلاً من "اختر رسالة"
        _add_heading_rtl(doc, "الرسالة", level=2)
        _add_paragraph_rtl(doc, str(message))
        doc.add_paragraph()  # سطر فارغ
        
        # الأهداف - استخدام "الأهداف" بدلاً من "الأهداف الاستراتيجية"
        _add_heading_rtl(doc, "الأهداف", level=2)
        for obj in strategy_data.get("objectives", []):
            _add_bullet_rtl(doc, str(obj))
        doc.add_paragraph()  # سطر فارغ
        
        # القيم - استخدام "القيم" بدلاً من "القيم الاستراتيجية"
        _add_heading_rtl(doc, "القيم", level=2)
        for val in strategy_data.get("values", []):
            _add_bullet_rtl(doc, str(val))

        # تأكيد RTL على كل الفقرات
        for p in doc.paragraphs:
            _rtl_paragraph(p)

        # حفظ في buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        st.success("✅ تم إنشاء ملف Word بنجاح!")
        st.download_button(
            label=get_text("export_to_word"),
            data=buffer,
            file_name=f"{strategy_data['name'].replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        st.error(f"خطأ في إنشاء ملف Word: {str(e)}")

def export_to_pdf(strategy_data):
    try:
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_RIGHT, TA_CENTER
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfbase import pdfmetrics
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import inch
        import arabic_reshaper
        from bidi.algorithm import get_display
        import os
        from PIL import Image as PILImage  # ✅ للحصول على الأبعاد الطبيعية
    except ImportError:
        st.error("المكتبات المطلوبة غير مثبتة. يرجى تثبيت: pip install reportlab arabic-reshaper python-bidi")
        return

    # دالة لمعالجة النص العربي
    def _process_arabic_text(text):
        try:
            if not text:
                return ""
            # إعادة تشكيل النص العربي
            reshaped_text = arabic_reshaper.reshape(str(text))
            # تطبيق الخوارزمية ثنائية الاتجاه
            bidi_text = get_display(reshaped_text)
            return bidi_text
        except:
            return str(text)

    # تسجيل خط Noto Naskh Arabic
    def _register_noto_font():
        try:
            # مسارات محتملة لخط Noto Naskh Arabic
            noto_paths = [
                "fonts/NotoNaskhArabic-Regular.ttf",
                "assets/fonts/NotoNaskhArabic-Regular.ttf",
                "NotoNaskhArabic-Regular.ttf",
                "C:/Windows/Fonts/NotoNaskhArabic-Regular.ttf",
                "./NotoNaskhArabic-Regular.ttf",
            ]
            
            for font_path in noto_paths:
                if os.path.exists(font_path):
                    # تسجيل الخط العادي
                    pdfmetrics.registerFont(TTFont('NotoNaskhArabic', font_path))
                    
                    # محاولة تسجيل الخط الغامق إذا كان متوفراً
                    bold_path = font_path.replace('Regular', 'Bold')
                    if os.path.exists(bold_path):
                        pdfmetrics.registerFont(TTFont('NotoNaskhArabic-Bold', bold_path))
                        return 'NotoNaskhArabic', 'NotoNaskhArabic-Bold'
                    else:
                        # استخدام نفس الخط للعادي والغامق
                        pdfmetrics.registerFont(TTFont('NotoNaskhArabic-Bold', font_path))
                        return 'NotoNaskhArabic', 'NotoNaskhArabic-Bold'
            
            # إذا لم يتم العثور على Noto Naskh Arabic
            st.warning("⚠️ خط Noto Naskh Arabic غير موجود. جاري البحث عن خطوط بديلة...")
            
            # البحث عن خطوط عربية بديلة
            arabic_fonts = [
                "arial.ttf", "tahoma.ttf", "times.ttf",
                "C:/Windows/Fonts/arial.ttf",
                "C:/Windows/Fonts/tahoma.ttf", 
            ]
            
            for font_path in arabic_fonts:
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont('ArabicFont', font_path))
                    pdfmetrics.registerFont(TTFont('ArabicFont-Bold', font_path))
                    st.warning(f"⚠️ استخدام خط بديل: {os.path.basename(font_path)}")
                    return 'ArabicFont', 'ArabicFont-Bold'
            
            st.error("❌ لم يتم العثور على أي خط يدعم العربية")
            return 'Helvetica', 'Helvetica-Bold'
            
        except Exception as e:
            st.error(f"❌ خطأ في تحميل الخط: {e}")
            return 'Helvetica', 'Helvetica-Bold'

    # إنشاء buffer للPDF
    buffer = BytesIO()
    
    # تسجيل خطوط Noto Naskh Arabic
    normal_font, bold_font = _register_noto_font()
    
    # إنشاء المستند
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=A4,
        rightMargin=40, 
        leftMargin=40, 
        topMargin=40, 
        bottomMargin=40,
        encoding='utf-8'
    )
    
    # استخدام القيم المحررة إذا كانت موجودة
    vision = strategy_data.get("edited_vision", strategy_data.get("vision", ""))
    message = strategy_data.get("edited_message", strategy_data.get("message", ""))
    
    # إنشاء الأنماط مع محاذاة عربية باستخدام Noto Naskh Arabic
    styles = getSampleStyleSheet()
    
    # أنماط مخصصة للعربية مع خط Noto Naskh Arabic
    style_title = ParagraphStyle(
        'ArabicTitle',
        parent=styles['Title'],
        fontName=bold_font,
        alignment=TA_CENTER,
        fontSize=20,
        leading=28,
        spaceAfter=20,
        wordWrap='RTL',
        textColor='#2c3e50'
    )
    
    style_heading = ParagraphStyle(
        'ArabicHeading',
        parent=styles['Heading1'],
        fontName=bold_font,
        alignment=TA_RIGHT,
        fontSize=16,
        spaceBefore=18,
        spaceAfter=10,
        wordWrap='RTL',
        textColor='#2980b9'
    )
    
    style_normal = ParagraphStyle(
        'ArabicNormal',
        parent=styles['Normal'],
        fontName=normal_font,
        alignment=TA_RIGHT,
        fontSize=14,
        leading=20,
        wordWrap='RTL',
        textColor='#2c3e50'
    )
    
    style_bullet = ParagraphStyle(
        'ArabicBullet',
        parent=styles['Normal'],
        fontName=normal_font,
        alignment=TA_RIGHT,
        fontSize=14,
        leading=18,
        leftIndent=30,
        wordWrap='RTL',
        textColor='#2c3e50'
    )
    
    # بناء المحتوى
    story = []
    
    # ✅ إضافة اللوجو بحجمه الطبيعي ومتوسطاً - مرة واحدة فقط
    logo_path = strategy_data.get("logo_path")
    if logo_path and os.path.exists(logo_path):
        try:
            # الحصول على الأبعاد الطبيعية للصورة
            with PILImage.open(logo_path) as img:
                width, height = img.size
            
            # حساب الحجم المناسب مع الحفاظ على النسبة
            max_width = 3.0 * inch  # أقصى عرض 3 إنش
            max_height = 2.5 * inch  # أقصى ارتفاع 2.5 إنش
            
            ratio = min(max_width/width, max_height/height)
            new_width = width * ratio
            new_height = height * ratio
            
            # إضافة الصورة بمحاذاة الوسط
            logo = Image(logo_path, width=new_width, height=new_height)
            logo.hAlign = 'CENTER'
            story.append(logo)
            story.append(Spacer(1, 20))
            
        except Exception as e:
            print(f"⚠️ Could not add logo to PDF with natural size: {e}")
            # محاولة بديلة بحجم افتراضي
            try:
                logo = Image(logo_path, width=2.5*inch, height=2.0*inch)
                logo.hAlign = 'CENTER'
                story.append(logo)
                story.append(Spacer(1, 15))
            except Exception as e2:
                print(f"⚠️ Could not add logo to PDF: {e2}")

    # العنوان الرئيسي
    title_text = _process_arabic_text(strategy_data["name"])
    if title_text:
        story.append(Paragraph(title_text, style_title))
        story.append(Spacer(1, 20))
    
    # الوصف إذا موجود
    if strategy_data.get("description"):
        desc_text = _process_arabic_text(strategy_data["description"])
        if desc_text:
            story.append(Paragraph(desc_text, style_normal))
            story.append(Spacer(1, 15))
    
    # الرؤية
    vision_heading = _process_arabic_text(get_text("select_vision"))
    vision_text = _process_arabic_text(vision)
    if vision_heading and vision_text:
        story.append(Paragraph(vision_heading, style_heading))
        story.append(Paragraph(vision_text, style_normal))
        story.append(Spacer(1, 15))
    
    # الرسالة
    message_heading = _process_arabic_text(get_text("select_message"))
    message_text = _process_arabic_text(message)
    if message_heading and message_text:
        story.append(Paragraph(message_heading, style_heading))
        story.append(Paragraph(message_text, style_normal))
        story.append(Spacer(1, 15))
    
    # الأهداف
    objectives_heading = _process_arabic_text(get_text("strategic_objectives"))
    if objectives_heading and strategy_data.get("objectives"):
        story.append(Paragraph(objectives_heading, style_heading))
        for obj in strategy_data["objectives"]:
            obj_text = _process_arabic_text("• " + str(obj))
            if obj_text:
                story.append(Paragraph(obj_text, style_bullet))
        story.append(Spacer(1, 15))
    
    # القيم
    values_heading = _process_arabic_text(get_text("strategic_values"))
    if values_heading and strategy_data.get("values"):
        story.append(Paragraph(values_heading, style_heading))
        for val in strategy_data["values"]:
            val_text = _process_arabic_text("• " + str(val))
            if val_text:
                story.append(Paragraph(val_text, style_bullet))
    
    # بناء PDF
    try:
        if story:
            doc.build(story)
            buffer.seek(0)
            
            st.success("✅ تم إنشاء ملف PDF بنجاح باستخدام خط Noto Naskh Arabic!")
            st.download_button(
                label=get_text("export_to_pdf"),
                data=buffer,
                file_name=f"{strategy_data['name'].replace(' ', '_')}.pdf",
                mime="application/pdf"
            )
        else:
            st.error("لا يوجد محتوى لتصديره إلى PDF")
    except Exception as e:
        st.error(f"خطأ في إنشاء PDF: {str(e)}")
        # محاولة استخدام نسخة مبسطة
        try:
            st.info("جاري استخدام نسخة مبسطة...")
            export_to_pdf_simple(strategy_data)
        except:
            st.error("فشل إنشاء PDF. يرجى التحقق من تثبيت المكتبات المطلوبة.")

# نسخة مبسطة للطوارئ
def export_to_pdf_simple(strategy_data):
    try:
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_RIGHT, TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import inch
    except ImportError:
        st.error("مكتبة reportlab غير مثبتة.")
        return

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    
    # أنماط مبسطة مع محاذاة عربية
    style_normal = ParagraphStyle(
        'SimpleArabic',
        parent=styles['Normal'],
        alignment=TA_RIGHT,
        fontSize=12,
        leading=16
    )
    
    style_heading = ParagraphStyle(
        'SimpleArabicHeading', 
        parent=styles['Heading2'],
        alignment=TA_RIGHT,
        fontSize=14,
        spaceAfter=8
    )
    
    story = []
    
    # ✅ إضافة اللوجو في النسخة المبسطة أيضاً - مرة واحدة فقط
    logo_path = strategy_data.get("logo_path")
    if logo_path and os.path.exists(logo_path):
        try:
            logo = Image(logo_path, width=2.0*inch, height=1.5*inch)
            logo.hAlign = 'CENTER'
            story.append(logo)
            story.append(Spacer(1, 15))
        except:
            pass
    
    vision = strategy_data.get("edited_vision", strategy_data.get("vision", ""))
    message = strategy_data.get("edited_message", strategy_data.get("message", ""))
    
    # إضافة المحتوى
    story.append(Paragraph(str(strategy_data["name"]), styles['Title']))
    story.append(Spacer(1, 12))
    
    if strategy_data.get("description"):
        story.append(Paragraph(str(strategy_data["description"]), style_normal))
        story.append(Spacer(1, 12))
    
    story.append(Paragraph("الرؤية:", style_heading))
    story.append(Paragraph(str(vision), style_normal))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("الرسالة:", style_heading))
    story.append(Paragraph(str(message), style_normal))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("الأهداف:", style_heading))
    for obj in strategy_data.get("objectives", []):
        story.append(Paragraph("• " + str(obj), style_normal))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("القيم:", style_heading))
    for val in strategy_data.get("values", []):
        story.append(Paragraph("• " + str(val), style_normal))
    
    doc.build(story)
    buffer.seek(0)
    
    st.download_button(
        label=get_text("export_to_pdf"),
        data=buffer,
        file_name=f"{strategy_data['name'].replace(' ', '_')}.pdf",
        mime="application/pdf"
    )

# --- Dialog لاختيار الرؤية ---
@st.dialog(get_text("select_vision"))
def vision_selector_dialog(username, current_data):
    visions = get_all_visions(username)
    if not visions:
        st.warning(get_text("no_visions_available"))
        return

    selected_vision = st.selectbox(
        get_text("select_vision_from_list"),
        options=visions,
        key="dialog_vision_select"
    )
    
    if st.button(get_text("confirm_selection")):
        current_data["selected_vision"] = selected_vision
        current_data["edited_vision"] = selected_vision
        st.session_state.vision_text_area = selected_vision
        st.rerun()

# --- Dialog لاختيار الرسالة ---
@st.dialog(get_text("select_message"))
def message_selector_dialog(username, current_data):
    original_vision = current_data.get("selected_vision")
    if not original_vision:
        st.error(get_text("please_select_vision_first"))
        return

    messages = get_messages_for_vision(original_vision, username)
    if not messages:
        st.warning(get_text("no_messages_available"))
        return

    selected_message = st.selectbox(
        get_text("select_message_from_list"),
        options=messages,
        key="dialog_message_select"
    )
    
    if st.button(get_text("confirm_selection")):
        current_data["selected_message"] = selected_message
        current_data["edited_message"] = selected_message
        st.session_state.message_text_area = selected_message
        st.rerun()

# --- وظيفة العرض الرئيسية ---
def render():
    st.title(get_text("strategy_management"))
    
    # === تطبيق الستايل على الـ Dialog ===
    lang = st.session_state.get("lang", "ar")
    is_rtl = lang == "ar"
    font_family = "'Tajawal', sans-serif" if lang == "ar" else "sans-serif"
    text_align = "right" if is_rtl else "left"
    direction = "rtl" if is_rtl else "ltr"
    
    st.markdown(f"""
    <style>
    /* تطبيق الستايل على جميع عناصر الـ Dialog */
    .stDialog * {{
        font-family: {font_family} !important;
        direction: {direction} !important;
        text-align: {text_align} !important;
    }}
    
    /* تطبيق الستايل على الزر */
    .stDialog .stButton > button {{
        border-radius: 12px !important;
        padding: 0.6rem 1.2rem !important;
        font-weight: 600 !important;
        font-size: 0.95rem !important;
        background: var(--brand) !important;
        color: white !important;
        border: none !important;
        box-shadow: 0 2px 6px rgba(0,0,0,0.15) !important;
        width: 100% !important;
    }}
    
    .stDialog .stButton > button:hover {{
        background: color-mix(in srgb, var(--brand), white 10%) !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 4px 12px rgba(0,0,0,0.25) !important;
    }}
    
    /* تطبيق الستايل على مربع النص (إذا وُجد في الـ Dialog) */
    .stDialog .stTextArea > div > div > textarea {{
        background: var(--input-bg) !important;
        color: var(--input-text) !important;
        border: 1px solid var(--input-border) !important;
        border-radius: 14px !important;
        padding: 0.75rem 1rem !important;
        font-size: 0.95rem !important;
        text-align: {text_align} !important;
        font-family: {font_family} !important;
    }}
    
    /* تطبيق الستايل على القائمة المنسدلة */
    .stDialog [data-baseweb="select"] {{
        direction: {direction} !important;
    }}
    
    .stDialog [data-baseweb="menu"] {{
        font-family: {font_family} !important;
        direction: {direction} !important;
    }}
    
    .stDialog [data-baseweb="menu"] li {{
        font-family: {font_family} !important;
        text-align: {text_align} !important;
    }}
    </style>
    """, unsafe_allow_html=True)
    
    user = get_current_user()
    if not user:
        st.error(get_text("must_be_logged_in"))
        return
    
    username = user["username"]
    
    if "strategy_builder_step" not in st.session_state:
        st.session_state.strategy_builder_step = 0
    if "current_strategy_data" not in st.session_state:
        st.session_state.current_strategy_data = {
            "name": "", "description": "", "vision": "", "message": "", "objectives": [], "values": [], "logo_path": None
        }
    if "editing_strategy_id" not in st.session_state:
        st.session_state.editing_strategy_id = None

    if st.session_state.strategy_builder_step == 0:
        render_strategy_list(username)
    else:
        render_strategy_builder(username)

# --- عرض قائمة الاستراتيجيات ---
def render_strategy_list(username):
    st.subheader(get_text("saved_strategies"))
    strategies = get_user_strategies(username)
    
    if not strategies:
        st.info(get_text("no_saved_strategies"))
    else:
        for strategy in strategies:
            with st.expander(f"📌 {strategy['name']}", expanded=False):
                # ✅ عرض اللوجو إذا موجود
                if strategy.get('logo_path') and os.path.exists(strategy['logo_path']):
                    st.image(strategy['logo_path'], width=100, caption="شعار الاستراتيجية")
                
                st.markdown(f"**{get_text('strategy_description')}**: {strategy['description'] or '—'}")
                st.markdown(f"**{get_text('select_vision')}**: {strategy['vision']}")
                st.markdown(f"**{get_text('select_message')}**: {strategy['message']}")
                st.markdown(f"**{get_text('strategic_objectives')}**: {', '.join(strategy['objectives'])}")
                st.markdown(f"**{get_text('strategic_values')}**: {', '.join(strategy['values'])}")
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    if st.button(get_text("edit"), key=f"edit_{strategy['id']}"):
                        st.session_state.current_strategy_data = {
                            "name": strategy["name"],
                            "description": strategy["description"],
                            "vision": strategy["vision"],
                            "message": strategy["message"],
                            "objectives": strategy["objectives"],
                            "values": strategy["values"],
                            "logo_path": strategy.get("logo_path"),
                            "edited_vision": strategy["vision"],
                            "edited_message": strategy["message"]
                        }
                        st.session_state.editing_strategy_id = strategy["id"]
                        st.session_state.strategy_builder_step = 1
                        st.rerun()
                with col2:
                    if st.button(get_text("delete"), key=f"delete_{strategy['id']}"):
                        if delete_strategy(strategy["id"], username):
                            # ✅ حذف ملف اللوجو أيضاً إذا كان موجوداً
                            if strategy.get('logo_path'):
                                delete_logo_file(strategy['logo_path'])
                            st.success(get_text("strategy_deleted_success"))
                            st.rerun()
                        else:
                            st.error(get_text("error_deleting_strategy"))
                with col3:
                    if st.button(get_text("view"), key=f"view_{strategy['id']}"):
                        st.session_state.view_strategy = strategy
                        st.session_state.strategy_builder_step = 5
                        st.session_state.editing_strategy_id = None
                        st.rerun()
    
    st.markdown("---")
    if st.button("➕ " + get_text("create_new_strategy")):
        st.session_state.current_strategy_data = {
            "name": "", "description": "", "vision": "", "message": "", "objectives": [], "values": [], "logo_path": None,
            "edited_vision": "", "edited_message": ""
        }
        st.session_state.editing_strategy_id = None
        st.session_state.strategy_builder_step = 1
        st.rerun()

# --- معالج بناء الاستراتيجية ---
def render_strategy_builder(username):
    steps = [
        get_text("select_vision"),
        get_text("select_message"),
        get_text("strategic_objectives"),
        get_text("strategic_values"),
        get_text("preview_and_save")
    ]
    step = st.session_state.strategy_builder_step - 1
    st.progress((step + 1) / 5, text=f"{steps[step]} ({step + 1}/5)")
    
    cols = st.columns(6)
    for i, label in enumerate(steps):
        disabled = (i < step - 1 or i > step + 1) and i != step
        if cols[i].button(label, disabled=disabled, use_container_width=True):
            st.session_state.strategy_builder_step = i + 1
            st.rerun()
    
    st.markdown("---")
    data = st.session_state.current_strategy_data

    # === الخطوة 1: الرؤية ===
    if st.session_state.strategy_builder_step == 1:
        # استخدام edited_vision إذا موجود، وإلا استخدام vision
        current_value = data.get("edited_vision", data.get("vision", ""))
        st.text_area(
            get_text("edit_or_create_vision"),
            value=current_value,
            height=120,
            key="vision_text_area"
        )

        col1, col2 = st.columns([2, 5])
        with col1:
            if st.button(get_text("browse_visions"), use_container_width=True):
                vision_selector_dialog(username, data)

        if st.button(get_text("save_and_continue"), use_container_width=True):
            final_vision = st.session_state.vision_text_area.strip()
            if not final_vision:
                st.error(get_text("vision_cannot_be_empty"))
            else:
                data["edited_vision"] = final_vision
                data["vision"] = final_vision  # تحديث الحقل الأصلي أيضاً
                if "selected_vision" not in data:
                    data["selected_vision"] = final_vision
                st.session_state.strategy_builder_step = 2
                st.rerun()

    # === الخطوة 2: الرسالة ===
    elif st.session_state.strategy_builder_step == 2:
        # استخدام edited_message إذا موجود، وإلا استخدام message
        current_value = data.get("edited_message", data.get("message", ""))
        st.text_area(
            get_text("edit_or_create_message"),
            value=current_value,
            height=120,
            key="message_text_area"
        )

        col1, col2 = st.columns([2, 5])
        with col1:
            if st.button(get_text("browse_messages"), use_container_width=True):
                message_selector_dialog(username, data)

        col3, col4 = st.columns(2)
        with col3:
            if st.button(get_text("back"), use_container_width=True):
                st.session_state.strategy_builder_step = 1
                st.rerun()
        with col4:
            if st.button(get_text("save_and_continue"), use_container_width=True):
                final_message = st.session_state.message_text_area.strip()
                if not final_message:
                    st.error(get_text("message_cannot_be_empty"))
                else:
                    data["edited_message"] = final_message
                    data["message"] = final_message  # تحديث الحقل الأصلي أيضاً
                    if "selected_message" not in data:
                        data["selected_message"] = final_message
                    st.session_state.strategy_builder_step = 3
                    st.rerun()

    # === الخطوة 3: الأهداف ===
    elif st.session_state.strategy_builder_step == 3:
        # جلب الأهداف المرتبطة بالرؤية والرسالة
        linked_objectives = get_objectives_for_vision_and_message(
            data.get("selected_vision", ""),
            data.get("selected_message", ""),
            username
        )
        all_objectives = get_all_objectives(username)
        
        if "selected_objectives_list" not in data:
            data["selected_objectives_list"] = data.get("objectives", [])
        
        st.subheader(get_text("strategic_objectives"))

        # القسم 1: الأهداف المرتبطة بالرسالة
        if linked_objectives:
            st.markdown(f"#### {get_text('objectives_linked_to_message')}")
            selected_from_linked = st.multiselect(
                get_text("select_from_linked_objectives"),
                options=linked_objectives,
                default=[obj for obj in linked_objectives if obj in data["selected_objectives_list"]],
                key="linked_objectives"
            )
            for obj in selected_from_linked:
                if obj not in data["selected_objectives_list"]:
                    data["selected_objectives_list"].append(obj)
        else:
            st.info(get_text("no_linked_objectives"))

        st.markdown("---")

        # القسم 2: جميع أهداف البنك
        st.markdown(f"#### {get_text('all_bank_objectives')}")
        selected_from_all = st.multiselect(
            get_text("select_from_all_objectives"),
            options=all_objectives,
            default=[obj for obj in all_objectives if obj in data["selected_objectives_list"]],
            key="all_objectives"
        )
        for obj in selected_from_all:
            if obj not in data["selected_objectives_list"]:
                data["selected_objectives_list"].append(obj)

        st.markdown("---")

        # القسم 3: إضافة هدف جديد
        new_objective = st.text_input(get_text("or_add_new_objective"))
        if st.button(get_text("add_objective")) and new_objective.strip():
            obj_clean = new_objective.strip()
            if obj_clean not in data["selected_objectives_list"]:
                data["selected_objectives_list"].append(obj_clean)
                save_custom_element(username, "objective", obj_clean)

        st.markdown("---")

        # القسم 4: الجدول التفاعلي
        if data["selected_objectives_list"]:
            st.markdown(f"#### {get_text('selected_objectives')}")
            for i, obj in enumerate(data["selected_objectives_list"]):
                col1, col2 = st.columns([5, 1])
                with col1:
                    st.text(obj)
                with col2:
                    if st.button("🗑️", key=f"del_obj_{i}"):
                        data["selected_objectives_list"].pop(i)
                        st.rerun()
        else:
            st.info(get_text("no_objectives_selected"))

        data["objectives"] = data["selected_objectives_list"]

        col1, col2 = st.columns(2)
        with col1:
            if st.button(get_text("back"), use_container_width=True):
                st.session_state.strategy_builder_step = 2
                st.rerun()
        with col2:
            if st.button(get_text("save_and_continue"), use_container_width=True):
                st.session_state.strategy_builder_step = 4
                st.rerun()

    # === الخطوة 4: القيم الاستراتيجية ===
    elif st.session_state.strategy_builder_step == 4:
        # جلب القيم المرتبطة بالرؤية والرسالة
        linked_values = get_values_for_vision_and_message(
            data.get("selected_vision", ""),
            data.get("selected_message", ""),
            username
        )
        all_values = get_all_values(username)
        
        if "selected_values_list" not in data:
            data["selected_values_list"] = data.get("values", [])
        
        st.subheader(get_text("strategic_values"))

        # القسم 1: القيم المرتبطة بالرسالة
        if linked_values:
            st.markdown(f"#### {get_text('values_linked_to_message')}")
            selected_from_linked = st.multiselect(
                get_text("select_from_linked_values"),
                options=linked_values,
                default=[val for val in linked_values if val in data["selected_values_list"]],
                key="linked_values"
            )
            for val in selected_from_linked:
                if val not in data["selected_values_list"]:
                    data["selected_values_list"].append(val)
        else:
            st.info(get_text("no_linked_values"))

        st.markdown("---")

        # القسم 2: جميع القيم من البنك
        st.markdown(f"#### {get_text('all_bank_values')}")
        selected_from_all = st.multiselect(
            get_text("select_from_all_values"),
            options=all_values,
            default=[val for val in all_values if val in data["selected_values_list"]],
            key="all_values"
        )
        for val in selected_from_all:
            if val not in data["selected_values_list"]:
                data["selected_values_list"].append(val)

        st.markdown("---")

        # القسم 3: إضافة قيمة جديدة
        new_value = st.text_input(get_text("or_add_new_value"))
        if st.button(get_text("add_value")) and new_value.strip():
            val_clean = new_value.strip()
            if val_clean not in data["selected_values_list"]:
                data["selected_values_list"].append(val_clean)
                save_custom_element(username, "value", val_clean)

        st.markdown("---")

        # القسم 4: الجدول التفاعلي (العرض + الحذف)
        if data["selected_values_list"]:
            st.markdown(f"#### {get_text('selected_values')}")
            for i, val in enumerate(data["selected_values_list"]):
                col1, col2 = st.columns([5, 1])
                with col1:
                    st.text(val)
                with col2:
                    if st.button("🗑️", key=f"del_val_{i}"):
                        data["selected_values_list"].pop(i)
                        st.rerun()
        else:
            st.info(get_text("no_values_selected"))

        # حفظ القائمة النهائية
        data["values"] = data["selected_values_list"]

        col1, col2 = st.columns(2)
        with col1:
            if st.button(get_text("back"), use_container_width=True):
                st.session_state.strategy_builder_step = 3
                st.rerun()
        with col2:
            if st.button(get_text("save_and_continue"), use_container_width=True):
                st.session_state.strategy_builder_step = 5
                st.rerun()

    # === الخطوة 5: المعاينة والحفظ ===
    elif st.session_state.strategy_builder_step == 5:
        st.subheader(get_text("preview_and_save"))
        
        # قسم اسم الاستراتيجية والوصف
        name = st.text_input(get_text("strategy_name"), value=data["name"])
        description = st.text_area(get_text("strategy_description"), value=data["description"])
        data["name"] = name
        data["description"] = description
        
        # ✅ القسم الجديد: رفع اللوجو
        st.markdown("---")
        st.subheader("📸 " + get_text("upload_logo"))
        
        col1, col2 = st.columns([2, 3])
        
        with col1:
            # عرض اللوجو الحالي إذا موجود
            current_logo_path = data.get("logo_path")
            if current_logo_path and os.path.exists(current_logo_path):
                st.image(current_logo_path, width=150, caption=get_text("current_logo"))
                if st.button("🗑️ " + get_text("remove_logo"), use_container_width=True):
                    if delete_logo_file(current_logo_path):
                        data["logo_path"] = None
                        st.success(get_text("logo_removed_success"))
                        st.rerun()
            else:
                st.info(get_text("no_logo_uploaded"))
        
        with col2:
            # رفع ملف جديد
            uploaded_logo = st.file_uploader(
                get_text("upload_logo"),
                type=["png", "jpg", "jpeg"],
                help="الحد الأقصى 5MB - PNG, JPG, JPEG"
            )
            
            if uploaded_logo:
                try:
                    # التحقق من الملف
                    validate_logo_file(uploaded_logo)
                    
                    # حفظ الملف مؤقتاً للعرض
                    temp_logo_path = save_logo_file(uploaded_logo, username, "temp")
                    if temp_logo_path:
                        st.image(temp_logo_path, width=150)
                        data["temp_logo_path"] = temp_logo_path
                        data["uploaded_logo"] = uploaded_logo
                        
                except Exception as e:
                    st.error(str(e))
        
        # استخدام القيم المحررة إذا كانت موجودة
        vision = data.get("edited_vision", data.get("vision", ""))
        message = data.get("edited_message", data.get("message", ""))
        
        # عرض معاينة الاستراتيجية
        st.markdown("---")
        st.subheader("👁️ " + get_text("preview_and_save"))
        
        # عرض اللوجو في المعاينة إذا موجود
        preview_logo_path = data.get("temp_logo_path") or data.get("logo_path")
        if preview_logo_path and os.path.exists(preview_logo_path):
            st.image(preview_logo_path, width=120)
        
        st.markdown(f"### {get_text('select_vision')}")
        st.markdown(f"> {vision}")
        st.markdown(f"### {get_text('select_message')}")
        st.markdown(f"> {message}")
        st.markdown(f"### {get_text('strategic_objectives')}")
        for obj in data["objectives"]:
            st.markdown(f"- {obj}")
        st.markdown(f"### {get_text('strategic_values')}")
        for val in data["values"]:
            st.markdown(f"- {val}")
        
        # أزرار التحكم
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            if st.button(get_text("back"), use_container_width=True):
                st.session_state.strategy_builder_step = 4
                st.rerun()
        with col2:
            if st.button(get_text("save_strategy"), use_container_width=True):
                if not name.strip():
                    st.error(get_text("strategy_name_required"))
                else:
                    # حفظ اللوجو إذا تم رفع واحد جديد
                    final_logo_path = data.get("logo_path")
                    if "uploaded_logo" in data:
                        # حذف اللوجو القديم إذا موجود
                        if data.get("logo_path"):
                            delete_logo_file(data["logo_path"])
                        
                        # حفظ اللوجو الجديد
                        strategy_id = st.session_state.editing_strategy_id or "new"
                        final_logo_path = save_logo_file(data["uploaded_logo"], username, strategy_id)
                    
                    # استخدام القيم المحررة إذا كانت موجودة
                    final_vision = data.get("edited_vision", data.get("vision", ""))
                    final_message = data.get("edited_message", data.get("message", ""))
                    
                    if st.session_state.editing_strategy_id:
                        success = update_strategy(
                            st.session_state.editing_strategy_id, username, name, description,
                            final_vision, final_message, data["objectives"], data["values"], final_logo_path
                        )
                    else:
                        strategy_id = create_strategy(
                            username, name, description,
                            final_vision, final_message, data["objectives"], data["values"], final_logo_path
                        )
                        success = True
                    
                    if success:
                        st.success(get_text("strategy_saved_success"))
                        st.session_state.strategy_builder_step = 0
                        st.rerun()
        with col3:
            if st.button(get_text("export_to_word"), use_container_width=True):
                export_to_word({
                    "name": name,
                    "description": description,
                    "vision": vision,
                    "message": message,
                    "objectives": data["objectives"],
                    "values": data["values"],
                    "edited_vision": vision,
                    "edited_message": message,
                    "logo_path": preview_logo_path  # ✅ إضافة اللوجو
                })
        with col4:
            if st.button(get_text("export_to_pdf"), use_container_width=True):
                export_to_pdf({
                    "name": name,
                    "description": description,
                    "vision": vision,
                    "message": message,
                    "objectives": data["objectives"],
                    "values": data["values"],
                    "edited_vision": vision,
                    "edited_message": message,
                    "logo_path": preview_logo_path  # ✅ إضافة اللوجو
                })
    
    st.markdown("---")
    if st.button("← " + get_text("back_to_strategy_list"), use_container_width=True):
        st.session_state.strategy_builder_step = 0
        st.rerun()